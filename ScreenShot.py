import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import openpyxl
import pyautogui
import io
import win32clipboard
from docx.shared import Inches
from openpyxl.drawing.image import Image as ExcelImage
from docx import Document
import tkinter.simpledialog


class ScreenshotApp:

    def __init__(self, root):

        self.snip_canvas = None
        self.snip_start_y = None
        self.snip_start_x = None
        self.snip_rect = None
        self.selection_rect = None
        self.snip_window = None
        self.root = root
        self.root.title("SnapIt")
        self.root.attributes('-topmost', True)  # Make the window stay on top
        self.root.resizable(False, False)  # Disable maximize window

        try:
            self.root.iconbitmap("Icon/SnapIt.ico")
        except Exception as e:
            print(f"Error loading icon: {e}")

        self.screenshots = []

        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

        # Snip button
        self.snip_btn = tk.Button(root, text="Snip", command=self.snip_screenshot)
        self.snip_btn.grid(row=0, column=0, padx=5, pady=5)

        # Status label
        self.status_label = tk.Label(root, text="Status: Ready.", anchor='w')
        self.status_label.grid(row=0, column=1, columnspan=5, pady=5, sticky='w')

        # Screenshot button
        self.screenshot_btn = tk.Button(root, text="Screenshot", command=self.take_screenshot)
        self.screenshot_btn.grid(row=1, column=0, padx=5, pady=5)

        # Copy last Screenshot
        self.copy_last_btn = tk.Button(root, text="Copy Last", command=self.copy_last_screenshot)
        self.copy_last_btn.grid(row=1, column=1, padx=5, pady=5)

        # Delete last Screenshot
        self.delete_last_btn = tk.Button(root, text="Delete last", command=self.delete_last_screenshot)
        self.delete_last_btn.grid(row=1, column=2, padx=5, pady=5)

        # Save dropdown
        self.save_option = tk.StringVar(root)
        self.save_option.set("Save to Excel")  # default value
        self.save_menu = ttk.Combobox(root, textvariable=self.save_option,
                                      values=["Save to Excel", "Save to Existing Excel", "Save to Word",
                                              "Save to Existing Word"], state='readonly')
        self.save_menu.grid(row=1, column=3, padx=5, pady=5)

        # Save button
        self.save_btn = tk.Button(root, text="Save", command=self.save_screenshot)
        self.save_btn.grid(row=1, column=4, padx=5, pady=5)

        # Reset button
        self.reset_btn = tk.Button(root, text="Reset", command=self.reset_screenshots)
        self.reset_btn.grid(row=1, column=5, padx=5, pady=5)

    def take_screenshot(self):

        # Temporarily hide the window
        self.root.attributes('-alpha', 0)  # Make the window completely transparent
        # Take the screenshot
        screenshot = pyautogui.screenshot()
        # Restore the window
        self.root.attributes('-alpha', 1)  # Restore visibility

        # Store the screenshot
        self.screenshots.append(screenshot)
        self.status_label.config(text=f"Status: Screenshot {len(self.screenshots)} taken")

    def snip_screenshot(self):
        self.root.withdraw()
        self.snip_window = tk.Toplevel(self.root)
        self.snip_window.attributes('-fullscreen', True)
        self.snip_window.attributes('-alpha', 0.3)
        self.snip_window.config(bg='black')
        self.snip_canvas = tk.Canvas(self.snip_window, cursor="cross", bg='black', highlightthickness=0)
        self.snip_canvas.pack(fill=tk.BOTH, expand=tk.YES)
        self.snip_canvas.bind('<Button-1>', self.start_snip)
        self.snip_canvas.bind('<B1-Motion>', self.update_snip)
        self.snip_canvas.bind('<ButtonRelease-1>', self.end_snip)
        self.selection_rect = None

    def start_snip(self, event):
        self.snip_start_x = event.x
        self.snip_start_y = event.y
        self.selection_rect = self.snip_canvas.create_rectangle(self.snip_start_x, self.snip_start_y, event.x, event.y,
                                                                outline='yellow', width=4)

    def update_snip(self, event):
        self.snip_canvas.coords(self.selection_rect, self.snip_start_x, self.snip_start_y, event.x, event.y)

    def end_snip(self, event):
        x1 = self.snip_start_x
        y1 = self.snip_start_y
        x2 = event.x
        y2 = event.y
        self.snip_window.destroy()
        self.root.deiconify()
        self.root.attributes('-alpha', 0)
        screenshot = pyautogui.screenshot(region=(x1, y1, x2 - x1, y2 - y1))
        self.root.attributes('-alpha', 1)
        self.screenshots.append(screenshot)
        self.status_label.config(text=f"Status: Snip {len(self.screenshots)}.")

    def save_screenshot(self):
        option = self.save_option.get()
        if option == "Save to Excel":
            self.save_to_excel()
        elif option == "Save to Existing Excel":
            self.save_to_existing_excel()
        elif option == "Save to Word":
            self.save_to_word()
        elif option == "Save to Existing Word":
            self.save_to_existing_word()

    def save_to_excel(self):
        if not self.screenshots:
            self.status_label.config(text="Status: No screenshots to save.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                                 filetypes=[("Excel files", "*.xlsx")])
        if not file_path:
            return

        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Screenshots"
            row_offset = 4

            for idx, screenshot in enumerate(self.screenshots):
                img_io = io.BytesIO()
                screenshot.save(img_io, format="PNG")
                img_io.seek(0)
                img = ExcelImage(img_io)
                img.anchor = f"A{row_offset}"
                ws.add_image(img)
                ws.row_dimensions[row_offset].height = screenshot.height // 15
                ws.column_dimensions['A'].width = screenshot.width // 15
                row_offset += int(screenshot.height / 15)

            wb.save(file_path)
            self.status_label.config(text=f"Status: Excel file saved!")
        except PermissionError:
            messagebox.showerror("Error", "The Excel file is open. Please close it and try again.")

    def save_to_existing_excel(self):
        if not self.screenshots:
            self.status_label.config(text="No screenshots to save.")
            return

        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if not file_path:
            return

        sheet_name = tk.simpledialog.askstring("Input", "Enter the sheet name:")

        if not sheet_name:
            self.status_label.config(text="Status: No sheet name provided.")
            return

        try:
            try:
                wb = openpyxl.load_workbook(file_path)
            except Exception as e:
                self.status_label.config(text="Failed to load the Excel file. Possibly File is open.")
                return

            ws = wb.create_sheet(title=sheet_name)

            row_offset = 4
            for idx, screenshot in enumerate(self.screenshots):
                img_io = io.BytesIO()
                screenshot.save(img_io, format="PNG")
                img_io.seek(0)
                img = ExcelImage(img_io)
                img.anchor = f"A{row_offset}"
                ws.add_image(img)
                ws.row_dimensions[row_offset].height = screenshot.height // 15
                ws.column_dimensions['A'].width = screenshot.width // 15
                row_offset += int(screenshot.height / 15)

            wb.save(file_path)
            self.status_label.config(text=f"Screenshots saved to existing {sheet_name} in {file_path}!")
        except PermissionError:
            messagebox.showerror("Error", "The Excel file is open. Please close it and try again.")

    def save_to_word(self):
        if not self.screenshots:
            self.status_label.config(text="Status: No screenshots to save.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word files", "*.docx")])
        if not file_path:
            return

        doc = Document()
        heading = doc.add_paragraph()
        run = doc.add_heading("1st Run Screenshots:")
        run.bold = True
        heading.style = 'Heading 1'
        doc.add_paragraph("")

        try:
            doc = Document()
            for idx, screenshot in enumerate(self.screenshots):
                img_io = io.BytesIO()
                screenshot.save(img_io, format="PNG")
                img_io.seek(0)
                doc.add_picture(img_io, width=Inches(6))
                img_io.seek(0)  # Reset the IO stream for Word
                doc.add_paragraph()  # Add space between screenshots

            doc.save(file_path)
            self.status_label.config(text=f"Status: Word file saved!")
        except PermissionError:
            messagebox.showerror("Error", "The Word file is open. Please close it and try again.")

    def save_to_existing_word(self):
        if not self.screenshots:
            self.status_label.config(text="Status: No screenshots to save.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word files", "*.docx")])
        if not file_path:
            return

        try:
            doc = Document()
            doc.add_page_break()
            heading = doc.add_paragraph()
            run = doc.add_heading("2nd Run Screenshots:")
            run.bold = True
            heading.style = 'Heading 1'
            doc.add_paragraph("")

            for idx, screenshot in enumerate(self.screenshots):
                img_io = io.BytesIO()
                screenshot.save(img_io, format="PNG")
                img_io.seek(0)
                doc.add_picture(img_io, width=Inches(6))
                doc.add_paragraph()  # Add space between screenshots

            doc.save(file_path)
            self.status_label.config(text=f"Status: Word file saved to existing Word.")

        except PermissionError:
            messagebox.showerror("Error", "The Word file is open. Please close it and try again.")
        except Exception as e:
            messagebox.showerror("Error", f"The Word file is open. Please close it and try again. {e}")

    def delete_last_screenshot(self):
        if self.screenshots:
            self.screenshots.pop()
            self.status_label.config(text=f"Status: Last Screenshot Deleted. {len(self.screenshots)} remaining.")
        else:
            self.status_label.config(text="Status: No Screenshots to Delete.")

    def copy_last_screenshot(self):
        if self.screenshots:
            last_screenshot = self.screenshots[-1]
            output = io.BytesIO()
            last_screenshot.save(output, format="BMP")
            data = output.getvalue()[14:]  # BMP have 14 byte header

            win32clipboard.OpenClipboard()
            win32clipboard.EmptyClipboard()
            win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)
            win32clipboard.CloseClipboard()

            self.status_label.config(text="Status: Last Screenshot copied to clipboard.")
        else:
            self.status_label.config(text="Status: No screenshot to copy.")

    def reset_screenshots(self):
        if self.screenshots:
            confirm = messagebox.askyesno("Confirm Reset", "Are you sure you want to reset all screenshots?")
            if not confirm:
                return
        self.screenshots = []  # Clear all stored screenshots
        self.status_label.config(text="Status: Screenshots reset.")

    def on_close(self):
        if self.screenshots:
            confirm = messagebox.askyesno("Confirm Exit", "Screenshots exist. Are you sure you want to exit?")
            if not confirm:
                return
        self.root.destroy()


if __name__ == "__main__":
    root = tk.Tk()
    app = ScreenshotApp(root)
    root.mainloop()
